from pathlib import Path

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader


def load_text(path: Path, content_type: str) -> list[dict]:
    suffix = path.suffix.lower()
    if suffix == ".pdf" or content_type == "application/pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    return [{"text": path.read_text(encoding="utf-8", errors="ignore"), "page": None}]


def _load_docx(path: Path) -> list[dict]:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return [{"text": "\n".join(parts), "page": None}]


def _load_pdf(path: Path) -> list[dict]:
    pages: list[dict] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"text": text, "page": index})
    if pages:
        return pages
    reader = PdfReader(str(path))
    return [{"text": page.extract_text() or "", "page": index} for index, page in enumerate(reader.pages, start=1)]
